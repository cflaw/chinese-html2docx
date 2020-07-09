import os
import subprocess
import re
import win32com.client
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor
from docx.shared import Pt
from docx.shared import Cm
from selenium import webdriver
from urllib.request import urlopen
from selenium.webdriver.support.ui import WebDriverWait
from selenium.common.exceptions import TimeoutException
from selenium.webdriver.firefox.options import Options
from datetime import datetime


def sorted_alphanumeric(data):
    def convert(text): return int(text) if text.isdigit() else text.lower()
    def alphanum(key): return [convert(c) for c in re.split('([0-9]+)', key)]
    return sorted(data, key=alphanum)


def getHTMLsimple(url):
    mainHtml = urlopen(url)
    mainContent = mainHtml.read().decode('gb2312', 'ignore')
    return mainContent


def getLatestChapter(mainContent):
    # latestChpt = re.search(r'最新章节：(.*?)</p>', mainContent).group(1)
    # chptNum = re.search(r'>第(.*?)章<', latestChpt).group(1)

    latestChpt = re.search(r'最新章节：(.*?)<\/a>', mainContent).group(1)
    chptNum = re.search(r'>第(.*?)章', latestChpt).group(1)
    return chptNum


def getChapters(dir, url):
    mainContent = getHTMLsimple(url)
    latest = getLatestChapter(mainContent)
    doneTill = 1
    links = {}

    for filename in sorted_alphanumeric(os.listdir(dir)):
        if filename.endswith('.completed'):
            doneTill = os.path.splitext(filename)[0]

    latest = int(latest)
    doneTill = int(doneTill)
    iterations = latest - doneTill
    for x in range(iterations):
        doneTill += 1
        regex = r'(<dd>|<li>)<a href( )?="(.*?)">第' + str(doneTill) + '章'
        chptUrl = re.search(regex, mainContent)
        if chptUrl:
            links[doneTill] = chptUrl.group(3)
    return links


def updateLatestChapter(directory, latest):
    for filename in os.listdir(directory):
        if filename.endswith('.completed'):
            full_filename = os.path.join(directory, filename)
            full_latestChapter = os.path.join(directory, str(latest))
            os.rename(full_filename, full_latestChapter)


def cleanHTML(content):
    regex1 = '\r'
    regex2 = '\n'
    regex3 = '\t'
    regex4 = '(&nbsp;){1,10}第([0-9]{1,4})章.{1,15}<br( )?(/)?>'
    # regex5 = '<br/><br/>'
    # replace_txt5 = '<br/>'

    cleaned_txt = content.replace(regex1, '')
    cleaned_txt = cleaned_txt.replace(regex2, '')
    cleaned_txt = cleaned_txt.replace(regex3, '')
    cleaned_txt = re.sub(regex4, '', cleaned_txt, flags=re.DOTALL)
    # cleaned_txt = cleaned_txt.replace(regex5, replace_txt5)

    return cleaned_txt


def removeNonsense(content):
    useless_words = [
        "更新最快",
        "78中文首发",
        "叶辰萧初然来源：",
        "叶辰萧初然",
        "首发",
        "手机端一秒住槟提供精彩\\小fx。",
        "&nbsp;??"
    ]

    for index, word in enumerate(useless_words):
        content = content.replace(useless_words[index], "")

    return content


def update_toc(docx_file):
    word = win32com.client.DispatchEx("Word.Application")
    doc = word.Documents.Open(docx_file)
    doc.TablesOfContents(1).Update()
    doc.Close(SaveChanges=True)
    word.Quit()


def waitForAjax(driver):
    try:
        WebDriverWait(driver, 5).until(
            lambda driver: driver.execute_script("return jQuery.active") == 0
        )
    except TimeoutException:
        raise TimeoutException("Wait for Ajax timed out.")


def convert2mobi():
    try:
        print(datetime.now(), ' - Converting to mobi.')
        FNULL = open(os.devnull, 'w')
        subprocess.call(
            [
                "ebook-convert",
                fullpath_combineDName,
                fullpath_combineFName,
                '--cover=' + fullpath_cover,
                '--authors=叶公子',
                '--title=叶辰萧初然小说'
            ], stdout=FNULL, stderr=subprocess.STDOUT)
        print(datetime.now(), ' - Converted to mobi.')
    except Exception as e:
        print(e)


# init directory and website location
directory = 'D:/reading/yechen/'

# old source
# base_url = 'https://www.biqupa.com'
# topic = '7_7817'

# new source
base_url = 'https://www.rzlib.net'
topic = 'b/23/23036'
url = '%s/%s/' % (base_url, topic)

combinedDocName = 'YeChenXiaoChuRanXiaoShuo.docx'
fullpath_combineDName = directory + combinedDocName

# included conversion to mobi directly in python script
combinedDocName_final = 'YeChenXiaoChuRanXiaoShuo.mobi'
fullpath_combineFName = directory + combinedDocName_final

# added cover page during mobi conversion process
fullpath_cover = directory + 'cover.jpg'

# init document creation and start formatting
document = Document(fullpath_combineDName)
firstPagebreak = True  # set flag for pagebreak

# formatting for whole document

# removed formatting as cover image is removed
# section = document.sections[0]
# section.left_margin = Cm(2.25)
# section.top_margin = Cm(0)

section = document.sections[0]
section.left_margin = Cm(2.25)
section.top_margin = Cm(2.0)

# formatting for heading of the chapter
chptValid = False

list_styles = [s for s in document.styles if s.type == WD_STYLE_TYPE.PARAGRAPH]
for style in list_styles:
    if(style.name == 'chpt'):
        chptValid = True

if (chptValid is False):
    styles = document.styles
    new_heading_style = styles.add_style('chpt', WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles['Heading 1']
    font = new_heading_style.font
    font.name = 'Calibri Light'
    font.size = Pt(12)
    font.bold = False
    font.color.rgb = RGBColor(0, 0, 0)

# check webpage for latest chapter and get array of links
chapters = getChapters(directory, url)

# debugging purposes
# chapters = {}
# chapters[1] = "/b/23/23036/46867544.html"

if (len(chapters) > 0):
    print(datetime.now(), ' - Found', len(chapters), 'new chapters.')

    options = Options()
    options.add_argument('-headless')
    wd = webdriver.Firefox(executable_path='geckodriver', options=options)
    # wd = webdriver.Firefox()

    for chapter in chapters:
        chapterURL = '%s/%s' % (base_url, chapters[chapter])
        # content = getHTML(chapterURL)

        pagebreak = document.add_paragraph().add_run()
        pagebreak.add_break(WD_BREAK.PAGE)

        # Add heading to the top of the page
        heading = document.add_paragraph("第"+str(chapter)+"章", style='chpt')

        # formatting for content
        heading.paragraph_format.space_before = Pt(0)
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.line_spacing = 1.5

        # get source from URL
        wd.get(chapterURL)
        waitForAjax(wd)

        # clean up html in preparations for processing
        source = cleanHTML(wd.page_source)
        source = removeNonsense(source)

        soup = BeautifulSoup(source, 'lxml')

        # get only the content
        content = soup.find("div", {"id": "chapter_content"})

        # add the cleaned up content into word doc
        run = paragraph.add_run(content.get_text(separator='\r\n'))
        run.font.name = 'Arial'
        run.font.size = Pt(10)

        # update to latest after completion
        updateLatestChapter(directory, str(chapter) + ".completed")

        if (chapter % 25 == 0 & bool(chapters)):
            document.save(fullpath_combineDName)

    wd.quit()

    if(bool(chapters)):
        document.save(fullpath_combineDName)
        update_toc(fullpath_combineDName)
        print(datetime.now(), ' - Updated', len(chapters), 'new chapters.')
        convert2mobi()
